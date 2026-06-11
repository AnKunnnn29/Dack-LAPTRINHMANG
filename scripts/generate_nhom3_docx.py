from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Nhom3_Report_Source.md"
OUTPUT = ROOT / "Nhom3_Report.docx"
FALLBACK_OUTPUT = ROOT / "Nhom3_Report_Final.docx"


def set_run_font(run, size: int = 13, bold: bool = False, italic: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.8)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(14 if style_name == "Heading 1" else 13)
        if style_name.startswith("Heading"):
            style.font.bold = True


def paragraph(doc: Document, text: str = "", align=None, bold: bool = False, italic: bool = False, size: int = 13, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 13, bold=True)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False, size: int = 12, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    lines = str(text).split("\n")
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            is_header = header and r_index == 0
            set_cell(cell, value, bold=is_header)
            if is_header:
                shade_cell(cell, "D9EAF7")
    paragraph(doc, "")


def add_placeholder(doc: Document, text: str) -> None:
    title = text.strip("[]")
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F2F2")
    set_cell(cell, f"[{title}]", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, title.replace("CHÈN HÌNH:", "Hình:"), align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F7F7")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.rstrip("\n").splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=10, name="Consolas")
    paragraph(doc, "")


def add_page_number(doc: Document) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trang ")
    set_run_font(run, size=11)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = p.add_run()
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_text)
    page_run._r.append(fld_end)


def add_cover(doc: Document) -> None:
    paragraph(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "TRƯỜNG ĐẠI HỌC SƯ PHẠM KỸ THUẬT TP. HCM", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "**********", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "")
    paragraph(doc, "BÁO CÁO ĐỒ ÁN CUỐI KÌ", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    paragraph(doc, "MÔN: LẬP TRÌNH MẠNG", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15)
    paragraph(doc, "")
    paragraph(doc, "XÂY DỰNG HỆ THỐNG NETWORK RECON + RISK PROFILER", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    paragraph(doc, "SỬ DỤNG MULTI-AGENT VÀ ML CHO AN TOÀN MẠNG", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    paragraph(doc, "")
    paragraph(doc, "GVHD: THẠC SĨ NGUYỄN ĐĂNG QUANG", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "MÃ HỌC PHẦN: NPRO430980", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "THỰC HIỆN: NHÓM 3", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "VŨ VĂN THÔNG - 23162098", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "NGUYỄN THÀNH AN - 23162001", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "")
    paragraph(doc, "TP. Hồ Chí Minh, Tháng 06 Năm 2026", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    doc.add_page_break()


def add_teacher_comment(doc: Document) -> None:
    paragraph(doc, "NHẬN XÉT CỦA GIÁO VIÊN", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15)
    for _ in range(15):
        paragraph(doc, "." * 138)
    paragraph(doc, "TP. Hồ Chí Minh, tháng 06 năm 2026", WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph(doc, "Giảng viên hướng dẫn", WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    paragraph(doc, "")
    paragraph(doc, "")
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    heading(doc, "BẢNG PHÂN CHIA NHIỆM VỤ NHÓM 3", 1)
    add_table(
        doc,
        [
            ["MSSV", "23162098", "23162001"],
            ["Tên thành viên", "Vũ Văn Thông", "Nguyễn Thành An"],
            [
                "Nhiệm vụ",
                "Pipeline & Recon Lead\n- Chính: .pi/tools/main_pipeline.py, common/tool_utils.py, recon/port_scanner.py, recon/dns_enum.py, recon/banner_grabber.py.\n- Vai trò: thiết kế safety gate, Stage 1 parallel recon, validate target/port, ghi JSON triage và kiểm thử pipeline offline.",
                "ML Risk & Report/Agentic Lead\n- Chính: risk/*.py, reporting/*.py, .pi/tools/pi_recon_agent.py, .pi/agents/*.md, .pi/chains/*.md, .pi/prompts/report_prompt.md.\n- Vai trò: feature engineering, Isolation Forest risk scoring, MITRE mapping, report fallback và Week 5 OpenAI tool-calling extension.",
            ],
            ["Mức độ hoàn thành", "Hoàn thành tốt nhiệm vụ 100%", "Hoàn thành tốt nhiệm vụ 100%"],
        ],
        header=False,
    )

    heading(doc, "MỤC LỤC", 1)
    for item in [
        "1. MỤC TIÊU ĐỀ TÀI",
        "2. GIỚI HẠN PHẠM VI",
        "3. NỘI DUNG THỰC HIỆN",
        "4. SỬ DỤNG AI ĐỂ GỢI Ý TOOLS",
        "5. CÁC TOOLS CẦN THIẾT SỬ DỤNG ĐỂ GIẢI QUYẾT VẤN ĐỀ NGOÀI NHỮNG TOOL DO AI ĐỀ XUẤT/GỢI Ý",
        "6. ĐÁNH GIÁ KẾT QUẢ THỰC HIỆN",
        "7. NỘI DUNG THỰC HIỆN VỚI MỤC ĐÍCH CẢI THIỆN HỆ THỐNG",
        "8. ĐÁNH GIÁ KẾT QUẢ THỰC HIỆN SAU KHI CẢI TIẾN",
        "9. BẢNG SO SÁNH PHIÊN BẢN ĐẦU VÀ SAU KHI CẢI TIẾN HỆ THỐNG",
        "10. KẾT LUẬN",
    ]:
        paragraph(doc, item)

    heading(doc, "DANH MỤC HÌNH ẢNH CẦN CHÈN", 1)
    for item in [
        "Hình 1. Sơ đồ kiến trúc tổng thể pipeline.",
        "Hình 2. Kết quả chạy main_pipeline.py trên terminal.",
        "Hình 3. Nội dung risk_profile.json hoặc ket_qua.md.",
        "Hình 4. Cấu trúc thư mục .pi/agents, .pi/skills, .pi/tools.",
        "Hình 5. Minh họa Week 5 agentic mode bằng pi_recon_agent.py.",
    ]:
        paragraph(doc, item)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    pending_paragraph: list[str] = []

    def flush_paragraph() -> None:
        nonlocal pending_paragraph
        if pending_paragraph:
            paragraph(doc, " ".join(item.strip() for item in pending_paragraph))
            pending_paragraph = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            block: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            add_code_block(doc, "\n".join(block))
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            heading(doc, stripped[2:].strip(), 1)
            index += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            heading(doc, stripped[3:].strip(), 2)
            index += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[2:].strip())
            set_run_font(run)
            index += 1
            continue

        if stripped.startswith("[CHÈN HÌNH:"):
            flush_paragraph()
            add_placeholder(doc, stripped)
            index += 1
            continue

        pending_paragraph.append(stripped)
        index += 1

    flush_paragraph()


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_page_number(doc)
    add_cover(doc)
    add_teacher_comment(doc)
    add_front_matter(doc)
    add_markdown_body(doc, SOURCE.read_text(encoding="utf-8"))
    try:
        doc.save(OUTPUT)
        print(OUTPUT)
    except PermissionError:
        doc.save(FALLBACK_OUTPUT)
        print(FALLBACK_OUTPUT)


if __name__ == "__main__":
    main()
